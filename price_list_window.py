import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from typing import List, Dict

class PriceListWindow:
    def __init__(self, parent, db):
        self.db = db
        self.selected_fans = []  # List of dicts: {'fan': fan_data, 'quantity': int}
        
        self.window = tk.Toplevel(parent)
        self.window.title("إنشاء عرض سعر / استفسار")
        self.window.geometry("1200x700")
        self.window.transient(parent)
        
        # Main container
        main_frame = ttk.Frame(self.window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text="مولد عرض السعر / الاستفسار", 
                               font=("Arial", 14, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 10))
        
        # Left panel - Search and available fans
        left_frame = ttk.LabelFrame(main_frame, text="المراوح المتاحة", padding="10")
        left_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(1, weight=1)
        
        # Search frame
        search_frame = ttk.Frame(left_frame)
        search_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        search_frame.columnconfigure(1, weight=1)
        
        ttk.Label(search_frame, text="بحث:").grid(row=0, column=0, padx=(0, 5))
        self.search_var = tk.StringVar()
        self.search_var.trace('w', self.on_search_change)
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var, width=25)
        search_entry.grid(row=0, column=1, sticky=(tk.W, tk.E))
        
        # Available fans treeview
        available_tree_frame = ttk.Frame(left_frame)
        available_tree_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        available_tree_frame.columnconfigure(0, weight=1)
        available_tree_frame.rowconfigure(0, weight=1)
        
        scrollbar_avail_y = ttk.Scrollbar(available_tree_frame, orient=tk.VERTICAL)
        scrollbar_avail_x = ttk.Scrollbar(available_tree_frame, orient=tk.HORIZONTAL)
        
        # Reverse column order for RTL display
        self.available_tree = ttk.Treeview(available_tree_frame,
                                           columns=("Qty", "Retail", "Wholesale", "Airflow", "Name"),
                                           show="headings",
                                           yscrollcommand=scrollbar_avail_y.set,
                                           xscrollcommand=scrollbar_avail_x.set,
                                           height=15)
        
        scrollbar_avail_y.config(command=self.available_tree.yview)
        scrollbar_avail_x.config(command=self.available_tree.xview)
        
        # Configure available tree columns (Arabic headings for fans, RTL order)
        self.available_tree.heading("Qty", text="كمية")
        self.available_tree.heading("Retail", text="مفرق")
        self.available_tree.heading("Wholesale", text="جملة")
        self.available_tree.heading("Airflow", text="غزارة")
        self.available_tree.heading("Name", text="نوع")
        
        self.available_tree.column("Qty", width=90, anchor=tk.E)
        self.available_tree.column("Retail", width=80, anchor=tk.E)
        self.available_tree.column("Wholesale", width=80, anchor=tk.E)
        self.available_tree.column("Airflow", width=100, anchor=tk.E)
        self.available_tree.column("Name", width=150, anchor=tk.E)
        
        self.available_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar_avail_y.grid(row=0, column=1, sticky=(tk.N, tk.S))
        scrollbar_avail_x.grid(row=1, column=0, sticky=(tk.W, tk.E))
        
        # Add button
        ttk.Button(left_frame, text="← إضافة إلى عرض السعر", 
                  command=self.add_to_price_list).grid(row=2, column=0, pady=10)
        
        # Right panel - Price list
        right_frame = ttk.LabelFrame(main_frame, text="عرض السعر / الاستفسار", padding="10")
        right_frame.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(1, weight=1)
        
        # Info label (price type is now per-item)
        info_frame = ttk.Frame(right_frame)
        info_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        ttk.Label(info_frame, text="انقر نقراً مزدوجاً على نوع السعر أو الكمية للتعديل", 
                 font=("Arial", 9)).pack()
        
        # Price list treeview
        price_list_tree_frame = ttk.Frame(right_frame)
        price_list_tree_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        price_list_tree_frame.columnconfigure(0, weight=1)
        price_list_tree_frame.rowconfigure(0, weight=1)
        
        scrollbar_list_y = ttk.Scrollbar(price_list_tree_frame, orient=tk.VERTICAL)
        scrollbar_list_x = ttk.Scrollbar(price_list_tree_frame, orient=tk.HORIZONTAL)
        
        # Reverse column order for RTL display
        self.price_list_tree = ttk.Treeview(price_list_tree_frame,
                                            columns=("Total", "Qty", "Price", "PriceType", "Airflow", "Name", "Order"),
                                            show="headings",
                                            yscrollcommand=scrollbar_list_y.set,
                                            xscrollcommand=scrollbar_list_x.set,
                                            height=15)
        
        scrollbar_list_y.config(command=self.price_list_tree.yview)
        scrollbar_list_x.config(command=self.price_list_tree.xview)
        
        # Configure price list tree columns (Arabic headings for fans, RTL order)
        self.price_list_tree.heading("Total", text="الإجمالي")
        self.price_list_tree.heading("Qty", text="كمية")
        self.price_list_tree.heading("Price", text="السعر")
        self.price_list_tree.heading("PriceType", text="نوع السعر")
        self.price_list_tree.heading("Airflow", text="غزارة")
        self.price_list_tree.heading("Name", text="نوع")
        self.price_list_tree.heading("Order", text="#")
        
        self.price_list_tree.column("Total", width=90, anchor=tk.E)
        self.price_list_tree.column("Qty", width=70, anchor=tk.E)
        self.price_list_tree.column("Price", width=90, anchor=tk.E)
        self.price_list_tree.column("PriceType", width=90, anchor=tk.E)
        self.price_list_tree.column("Airflow", width=120, anchor=tk.E)
        self.price_list_tree.column("Name", width=160, anchor=tk.E)
        self.price_list_tree.column("Order", width=40, anchor=tk.E)
        
        # Bind double-click to edit quantity or price type
        self.price_list_tree.bind('<Double-1>', self.on_item_double_click)
        
        self.price_list_tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar_list_y.grid(row=0, column=1, sticky=(tk.N, tk.S))
        scrollbar_list_x.grid(row=1, column=0, sticky=(tk.W, tk.E))
        
        # Order control frame
        order_frame = ttk.LabelFrame(right_frame, text="الترتيب", padding="5")
        order_frame.grid(row=2, column=0, pady=5, sticky=(tk.W, tk.E))
        
        ttk.Button(order_frame, text="↑ نقل لأعلى", 
                  command=self.move_item_up, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(order_frame, text="↓ نقل لأسفل", 
                  command=self.move_item_down, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(order_frame, text="ترتيب حسب الاسم ↑", 
                  command=lambda: self.sort_by_id(ascending=True), width=18).pack(side=tk.LEFT, padx=2)
        ttk.Button(order_frame, text="ترتيب حسب الاسم ↓", 
                  command=lambda: self.sort_by_id(ascending=False), width=18).pack(side=tk.LEFT, padx=2)
        
        # Buttons frame
        buttons_frame = ttk.Frame(right_frame)
        buttons_frame.grid(row=3, column=0, pady=10)
        
        ttk.Button(buttons_frame, text="حذف المحدد", 
                  command=self.remove_from_price_list).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="مسح الكل", 
                  command=self.clear_price_list).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="تصدير إلى Word", 
                  command=self.export_to_word).pack(side=tk.LEFT, padx=5)
        
        # Load initial data
        self.refresh_available_fans()
    
    def refresh_available_fans(self):
        """Refresh the available fans list"""
        # Clear existing items
        for item in self.available_tree.get_children():
            self.available_tree.delete(item)
        
        # Get all fans
        search_term = self.search_var.get().strip()
        if search_term:
            fans = self.db.search_fans(search_term)
        else:
            fans = self.db.get_all_fans()
        
        # Insert into treeview (use iid to store fan ID, RTL order)
        for fan in fans:
            self.available_tree.insert("", tk.END, iid=str(fan['id']), values=(
                fan['quantity'],
                f"${fan['price_retail']:.2f}",
                f"${fan['price_wholesale']:.2f}",
                fan.get('airflow') or "",
                fan['name']
            ))
    
    def on_search_change(self, *args):
        """Handle search input changes"""
        self.refresh_available_fans()
    
    def add_to_price_list(self):
        """Add selected fan to price list"""
        selection = self.available_tree.selection()
        if not selection:
            messagebox.showwarning("لا يوجد تحديد", "يرجى تحديد مروحة للإضافة.")
            return
        
        # Get fan ID from the treeview item identifier (iid)
        fan_id = int(selection[0])
        
        # Check if already in price list
        if any(item_data['fan']['id'] == fan_id for item_data in self.selected_fans):
            messagebox.showinfo("تمت الإضافة مسبقاً", "هذه المروحة موجودة بالفعل في عرض السعر.")
            return
        
        # Get fan data
        fan = self.db.get_fan_by_id(fan_id)
        if fan:
            # Ask user to select price type for this item
            price_type = self.select_price_type_dialog(fan['name'])
            if price_type is None:
                return  # User cancelled
            
            # Add with default quantity of 1, selected price type, and order index
            order_index = len(self.selected_fans) + 1
            self.selected_fans.append({
                'fan': fan, 
                'quantity': 1,
                'price_type': price_type,
                'order': order_index
            })
            self.update_price_list()
    
    def remove_from_price_list(self):
        """Remove selected fan from price list"""
        selection = self.price_list_tree.selection()
        if not selection:
            messagebox.showwarning("لا يوجد تحديد", "يرجى تحديد مروحة للحذف.")
            return
        
        item = self.price_list_tree.item(selection[0])
        fan_name = item['values'][5]  # Name is in column 5 (RTL order)
        
        # Remove from selected fans
        self.selected_fans = [f for f in self.selected_fans if f['fan']['name'] != fan_name]
        self.update_price_list()
    
    def select_price_type_dialog(self, fan_name):
        """Dialog to select price type for an item"""
        dialog = tk.Toplevel(self.window)
        dialog.title("اختر نوع السعر")
        dialog.geometry("300x150")
        dialog.transient(self.window)
        dialog.grab_set()
        
        # Center the dialog
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (dialog.winfo_width() // 2)
        y = (dialog.winfo_screenheight() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        result = {'price_type': None}
        
        ttk.Label(dialog, text=f"اختر نوع السعر لـ:\n{fan_name}", 
                 font=("Arial", 10)).pack(pady=10)
        
        price_type_var = tk.StringVar(value="retail")
        
        frame = ttk.Frame(dialog)
        frame.pack(pady=10)
        
        ttk.Radiobutton(frame, text="مفرق", variable=price_type_var, 
                       value="retail").pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(frame, text="جملة", variable=price_type_var, 
                       value="wholesale").pack(side=tk.LEFT, padx=10)
        
        def save():
            result['price_type'] = price_type_var.get()
            dialog.destroy()
        
        def cancel():
            dialog.destroy()
        
        button_frame = ttk.Frame(dialog)
        button_frame.pack(pady=10)
        ttk.Button(button_frame, text="موافق", command=save).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="إلغاء", command=cancel).pack(side=tk.LEFT, padx=5)
        
        self.window.wait_window(dialog)
        return result['price_type']
    
    def on_item_double_click(self, event):
        """Handle double-click on price list item"""
        selection = self.price_list_tree.selection()
        if not selection:
            return
        
        item = self.price_list_tree.item(selection[0])
        # In RTL order: Total=0, Qty=1, Price=2, PriceType=3, Airflow=4, Name=5, Order=6
        fan_name = item['values'][5]  # Name is in column 5
        column = self.price_list_tree.identify_column(event.x)
        
        # Find the item in selected_fans
        for item_data in self.selected_fans:
            if item_data['fan']['name'] == fan_name:
                # Column indices in RTL: Total=#1, Qty=#2, Price=#3, PriceType=#4, Airflow=#5, Name=#6, Order=#7
                # Values array: Total=0, Qty=1, Price=2, PriceType=3, Airflow=4, Name=5, Order=6
                if column == '#4':  # PriceType column
                    # Edit price type
                    new_price_type = self.select_price_type_dialog(fan_name)
                    if new_price_type is not None:
                        item_data['price_type'] = new_price_type
                        self.update_price_list()
                elif column == '#2':  # Quantity column (Qty is now #2 in RTL)
                    # Edit quantity
                    current_qty = item['values'][1]
                    new_qty = simpledialog.askinteger("تعديل الكمية", 
                                                      f"أدخل الكمية لـ {fan_name}:",
                                                      initialvalue=current_qty,
                                                      minvalue=1,
                                                      parent=self.window)
                    if new_qty is not None:
                        item_data['quantity'] = new_qty
                        self.update_price_list()
                break
    
    def clear_price_list(self):
        """Clear all fans from price list"""
        if self.selected_fans and messagebox.askyesno("تأكيد المسح", 
                                                      "هل تريد مسح جميع العناصر من عرض السعر؟"):
            self.selected_fans = []
            self.update_price_list()
    
    def update_price_list(self):
        """Update the price list display"""
        # Clear existing items
        for item in self.price_list_tree.get_children():
            self.price_list_tree.delete(item)
        
        # Reassign order numbers to maintain sequence
        for idx, item_data in enumerate(self.selected_fans, 1):
            item_data['order'] = idx
        
        # Insert selected fans with their individual price types
        for item_data in self.selected_fans:
            fan = item_data['fan']
            qty = item_data['quantity']
            order = item_data.get('order', 0)
            price_type = item_data.get('price_type', 'retail')  # Default to retail if not set
            price = fan['price_retail'] if price_type == 'retail' else fan['price_wholesale']
            total = price * qty
            price_type_label = "مفرق" if price_type == 'retail' else "جملة"
            # Insert in RTL order (reversed)
            self.price_list_tree.insert("", tk.END, values=(
                f"${total:.2f}",
                qty,
                f"${price:.2f}",
                price_type_label,
                fan.get('airflow') or "",
                fan['name'],
                order
            ))
    
    def move_item_up(self):
        """Move selected item up in the list"""
        selection = self.price_list_tree.selection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select an item to move.")
            return
        
        item = self.price_list_tree.item(selection[0])
        fan_name = item['values'][5]  # Name is in column 5 (RTL order: Total=0, Qty=1, Price=2, PriceType=3, Airflow=4, Name=5, Order=6)
        
        # Find the item index
        for idx, item_data in enumerate(self.selected_fans):
            if item_data['fan']['name'] == fan_name:
                if idx > 0:
                    # Swap with previous item
                    self.selected_fans[idx], self.selected_fans[idx - 1] = \
                        self.selected_fans[idx - 1], self.selected_fans[idx]
                    self.update_price_list()
                    # Reselect the moved item
                    children = self.price_list_tree.get_children()
                    if idx - 1 < len(children):
                        self.price_list_tree.selection_set(children[idx - 1])
                else:
                    messagebox.showinfo("معلومات", "العنصر موجود بالفعل في الأعلى.")
                break
    
    def move_item_down(self):
        """Move selected item down in the list"""
        selection = self.price_list_tree.selection()
        if not selection:
            messagebox.showwarning("لا يوجد تحديد", "يرجى تحديد عنصر للنقل.")
            return
        
        item = self.price_list_tree.item(selection[0])
        fan_name = item['values'][5]  # Name is in column 5 (RTL order)
        
        # Find the item index
        for idx, item_data in enumerate(self.selected_fans):
            if item_data['fan']['name'] == fan_name:
                if idx < len(self.selected_fans) - 1:
                    # Swap with next item
                    self.selected_fans[idx], self.selected_fans[idx + 1] = \
                        self.selected_fans[idx + 1], self.selected_fans[idx]
                    self.update_price_list()
                    # Reselect the moved item
                    children = self.price_list_tree.get_children()
                    if idx + 1 < len(children):
                        self.price_list_tree.selection_set(children[idx + 1])
                else:
                    messagebox.showinfo("معلومات", "العنصر موجود بالفعل في الأسفل.")
                break
    
    def sort_by_id(self, ascending=True):
        """Sort items by fan ID"""
        if not self.selected_fans:
            return
        
        self.selected_fans.sort(key=lambda x: x['fan']['id'], reverse=not ascending)
        self.update_price_list()
        messagebox.showinfo("تم الترتيب", f"تم ترتيب العناصر حسب الاسم ({'تصاعدي' if ascending else 'تنازلي'})")
    
    def export_to_word(self):
        """Export price list to a Word document using a template file"""
        if not self.selected_fans:
            messagebox.showwarning("قائمة فارغة", "عرض السعر فارغ.")
            return
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import os
        except ImportError:
            messagebox.showerror("Error", 
                "python-docx library is required. Please install it using:\n"
                "pip install python-docx")
            return
        
        from tkinter import filedialog
        from datetime import datetime
        
        # Look for template file
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
        except:
            current_dir = os.getcwd()
        
        template_paths = [
            "template.docx",
            "format.docx",
            os.path.join(current_dir, "template.docx"),
            os.path.join(current_dir, "format.docx")
        ]
        
        template_path = None
        for path in template_paths:
            if os.path.exists(path):
                template_path = path
                break
        
        if not template_path:
            # Ask user to select template file
            self.window.update()
            self.window.lift()
            self.window.focus_force()
            
            template_path = filedialog.askopenfilename(
                title="اختر ملف القالب (Template File)",
                filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
                parent=self.window
            )
            
            if not template_path:
                messagebox.showwarning("تحذير", "يجب اختيار ملف القالب للمتابعة.")
                return
        
        # Ensure window is active before showing dialogs
        self.window.update()
        self.window.lift()
        self.window.focus_force()
        
        # Get customer name (with parent window to ensure it stays on top)
        customer_name = simpledialog.askstring("اسم العميل", 
                                               "أدخل اسم العميل:",
                                               initialvalue="السيد نبيل حميدان المحترم",
                                               parent=self.window)
        if customer_name is None:
            return  # User cancelled
        
        # Ensure window is still active before next dialog
        self.window.update()
        self.window.lift()
        self.window.focus_force()
        
        # Get date (optional, defaults to today) - with parent window
        date_str = simpledialog.askstring("التاريخ", 
                                         "أدخل التاريخ (YYYY/MM/DD) أو اتركه فارغاً لليوم:",
                                         initialvalue=datetime.now().strftime("%Y/%m/%d"),
                                         parent=self.window)
        if date_str is None:
            return  # User cancelled
        if not date_str.strip():
            date_str = datetime.now().strftime("%Y/%m/%d")
        
        # Ensure window is raised before file dialog
        self.window.update()
        self.window.lift()
        self.window.focus_force()
        
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            title="حفظ عرض السعر",
            parent=self.window
        )
        
        if filename:
            try:
                # Load template document
                doc = Document(template_path)
                
                # Helper function to replace placeholders in text
                def replace_placeholder_in_runs(runs, placeholder, replacement):
                    """Replace placeholder text in paragraph runs"""
                    full_text = ''.join(run.text for run in runs)
                    if placeholder in full_text:
                        # Clear all runs
                        for run in runs:
                            run.text = ''
                        # Recreate with replacement
                        if runs:
                            runs[0].text = full_text.replace(placeholder, replacement)
                
                # Replace placeholders in all paragraphs
                placeholders = {
                    '{CUSTOMER_NAME}': customer_name,
                    '{DATE}': date_str,
                    '{{CUSTOMER_NAME}}': customer_name,
                    '{{DATE}}': date_str,
                }
                
                for para in doc.paragraphs:
                    for placeholder, replacement in placeholders.items():
                        if placeholder in para.text:
                            # Replace in all runs
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, replacement)
                            # Also check the paragraph text directly
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, replacement)
                
                # Replace placeholders in table cells
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for placeholder, replacement in placeholders.items():
                                    if placeholder in para.text:
                                        # Replace in all runs
                                        for run in para.runs:
                                            if placeholder in run.text:
                                                run.text = run.text.replace(placeholder, replacement)
                                        # Also check the paragraph text directly
                                        if placeholder in para.text:
                                            para.text = para.text.replace(placeholder, replacement)
                
                # Find and populate product table
                # Look for a table with placeholder {PRODUCTS_TABLE} or find the first suitable table
                product_table = None
                product_table_placeholder_found = False
                
                # First, try to find a table with {PRODUCTS_TABLE} placeholder
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if '{PRODUCTS_TABLE}' in para.text or '{{PRODUCTS_TABLE}}' in para.text:
                                    product_table = table
                                    product_table_placeholder_found = True
                                    # Clear the placeholder cell
                                    para.text = ''
                                    break
                        if product_table:
                            break
                    if product_table:
                        break
                
                # If no placeholder found, use the first table (assuming it's the product table)
                if not product_table and doc.tables:
                    product_table = doc.tables[0]
                
                if product_table:
                    # Clear existing rows except header (if placeholder was found, clear all)
                    if product_table_placeholder_found:
                        # Remove all rows
                        for i in range(len(product_table.rows) - 1, -1, -1):
                            product_table._element.remove(product_table.rows[i]._element)
                        # Add header row
                        header_row = product_table.add_row()
                    else:
                        # Keep first row as header, clear others
                        while len(product_table.rows) > 1:
                            product_table._element.remove(product_table.rows[-1]._element)
                        header_row = product_table.rows[0]
                    
                    # Set table RTL properties
                    tbl = product_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    tblPr.set(qn('w:bidiVisual'), '1')
                    tblPr.set(qn('w:jc'), 'right')
                    
                    # Ensure header exists with 4 columns
                    header_cells = header_row.cells
                    while len(header_cells) < 4:
                        header_row._element.append(OxmlElement('w:tc'))
                        header_cells = header_row.cells
                    
                    # Set header text (RTL order: الإجمالي, عدد, الإفرادي, النوع)
                    headers = ["الإجمالي", "عدد", "الإفرادي", "النوع"]
                    for i, header_text in enumerate(headers):
                        if i < len(header_cells):
                            cell = header_cells[i]
                            # Clear existing content
                            cell.text = ''
                            para = cell.paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            pPr = para._element.get_or_add_pPr()
                            pPr.set(qn('w:bidi'), '1')
                            run = para.add_run(header_text)
                            run.bold = True
                    
                    # Add product rows
                    grand_total = 0
                    for item_data in self.selected_fans:
                        fan = item_data['fan']
                        qty = item_data['quantity']
                        item_price_type = item_data.get('price_type', 'retail')
                        unit_price = fan['price_retail'] if item_price_type == 'retail' else fan['price_wholesale']
                        total_price = unit_price * qty
                        grand_total += total_price
                        
                        row_cells = product_table.add_row().cells
                        
                        # Column 0: الإجمالي (Total)
                        if len(row_cells) > 0:
                            total_cell = row_cells[0]
                            total_para = total_cell.paragraphs[0]
                            total_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            total_para.text = f"$ {total_price:.0f}"
                        
                        # Column 1: عدد (Quantity)
                        if len(row_cells) > 1:
                            qty_cell = row_cells[1]
                            qty_para = qty_cell.paragraphs[0]
                            qty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            qty_para.text = str(qty)
                        
                        # Column 2: الإفرادي (Unit Price)
                        if len(row_cells) > 2:
                            price_cell = row_cells[2]
                            price_para = price_cell.paragraphs[0]
                            price_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            price_para.text = f"{unit_price:.0f}"
                        
                        # Column 3: النوع (Type/Description)
                        if len(row_cells) > 3:
                            type_cell = row_cells[3]
                            type_cell.text = ''  # Clear cell
                            type_para = type_cell.paragraphs[0]
                            type_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            pPr = type_para._element.get_or_add_pPr()
                            pPr.set(qn('w:bidi'), '1')
                            
                            # Add description first (if exists), then name
                            description = fan.get('description')
                            description_text = description.strip() if description else ''
                            if description_text:
                                type_para.add_run(description_text)
                                type_para = type_cell.add_paragraph()
                                type_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                pPr = type_para._element.get_or_add_pPr()
                                pPr.set(qn('w:bidi'), '1')
                            
                            # Add name
                            type_para.add_run(fan['name'])
                            
                            # Add airflow if exists
                            airflow = fan.get('airflow')
                            if airflow:
                                type_para = type_cell.add_paragraph()
                                type_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                pPr = type_para._element.get_or_add_pPr()
                                pPr.set(qn('w:bidi'), '1')
                                type_para.add_run(f"Airflow: {airflow}")
                
                # Ensure RTL is set on all paragraphs and table cells
                for para in doc.paragraphs:
                    pPr = para._element.get_or_add_pPr()
                    pPr.set(qn('w:bidi'), '1')
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                pPr = para._element.get_or_add_pPr()
                                pPr.set(qn('w:bidi'), '1')
                
                # Save document
                doc.save(filename)
                
                messagebox.showinfo("نجح", f"تم تصدير عرض السعر إلى {filename}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export: {str(e)}")
                import traceback
                traceback.print_exc()

